"""文档处理 MCP Server

提供文档格式转换服务：
- docx 转 markdown
- 文档内容提取
- 文本格式转换
"""

import asyncio
import json
import os
import re
from datetime import datetime
from typing import Any, Dict, List, Optional

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from mcp.server import Server
from mcp.types import Tool, TextContent
from mcp.server.stdio import stdio_server


# 创建 MCP Server
server = Server("flood-agent-document")

# 配置
DATA_DIR = "./data"
OUTPUT_DIR = "./plans"


def _ensure_dirs():
    """确保目录存在"""
    os.makedirs(DATA_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)


def _validate_path(path: str, allowed_dirs: List[str]) -> str:
    """验证路径是否在允许范围内"""
    abs_path = os.path.abspath(path)
    for allowed in allowed_dirs:
        if abs_path.startswith(os.path.abspath(allowed)):
            return abs_path
    raise ValueError(f"路径不在允许范围内: {path}")


@server.list_tools()
async def list_tools() -> List[Tool]:
    """定义可用工具"""
    tools = [
        Tool(
            name="docx_to_markdown",
            description="将 docx 文件转换为 markdown 格式",
            inputSchema={
                "type": "object",
                "properties": {
                    "input_file": {
                        "type": "string",
                        "description": "输入 docx 文件路径（相对于 data 目录）"
                    },
                    "output_file": {
                        "type": "string",
                        "description": "输出 markdown 文件路径（相对于 plans 目录），可选"
                    },
                    "include_metadata": {
                        "type": "boolean",
                        "description": "是否包含文档元数据",
                        "default": True
                    },
                    "extract_images": {
                        "type": "boolean",
                        "description": "是否提取图片（当前版本忽略）",
                        "default": False
                    }
                },
                "required": ["input_file"]
            }
        ),
        Tool(
            name="read_docx_content",
            description="读取 docx 文件的纯文本内容",
            inputSchema={
                "type": "object",
                "properties": {
                    "input_file": {
                        "type": "string",
                        "description": "输入 docx 文件路径"
                    }
                },
                "required": ["input_file"]
            }
        ),
        Tool(
            name="list_documents",
            description="列出指定目录下的文档文件",
            inputSchema={
                "type": "object",
                "properties": {
                    "directory": {
                        "type": "string",
                        "description": "目录名称（data 或 plans）",
                        "default": "data"
                    },
                    "extension": {
                        "type": "string",
                        "description": "文件扩展名过滤，如 .docx",
                        "default": ".docx"
                    }
                }
            }
        ),
        Tool(
            name="create_sample_docx",
            description="创建示例 docx 文件用于测试",
            inputSchema={
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "文件名",
                        "default": "sample.docx"
                    },
                    "title": {
                        "type": "string",
                        "description": "文档标题",
                        "default": "示例文档"
                    }
                }
            }
        )
    ]
    return tools


@server.call_tool()
async def call_tool(name: str, arguments: Dict[str, Any]) -> List[TextContent]:
    """处理工具调用"""
    _ensure_dirs()
    
    try:
        if name == "docx_to_markdown":
            return await _handle_docx_to_markdown(arguments)
        elif name == "read_docx_content":
            return await _handle_read_docx_content(arguments)
        elif name == "list_documents":
            return await _handle_list_documents(arguments)
        elif name == "create_sample_docx":
            return await _handle_create_sample_docx(arguments)
        else:
            raise ValueError(f"未知工具: {name}")
    except Exception as e:
        return [TextContent(
            type="text",
            text=json.dumps({
                "success": False,
                "error": str(e),
                "error_type": type(e).__name__
            }, ensure_ascii=False)
        )]


async def _handle_docx_to_markdown(args: Dict[str, Any]) -> List[TextContent]:
    """处理 docx 转 markdown"""
    if not DOCX_AVAILABLE:
        return [TextContent(
            type="text",
            text=json.dumps({
                "success": False,
                "error": "python-docx 未安装，请运行: pip install python-docx"
            }, ensure_ascii=False)
        )]
    
    input_file = args["input_file"]
    output_file = args.get("output_file")
    include_metadata = args.get("include_metadata", True)
    
    # 构建完整路径
    input_path = os.path.join(DATA_DIR, input_file)
    input_path = _validate_path(input_path, [DATA_DIR])
    
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"文件不存在: {input_file}")
    
    # 读取 docx
    doc = Document(input_path)
    
    # 转换内容
    markdown_lines = []
    
    # 添加元数据
    if include_metadata:
        markdown_lines.append("---")
        markdown_lines.append(f"title: {os.path.splitext(input_file)[0]}")
        markdown_lines.append(f"converted_at: {datetime.now().isoformat()}")
        markdown_lines.append(f"source_file: {input_file}")
        markdown_lines.append("---")
        markdown_lines.append("")
    
    # 处理段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            markdown_lines.append("")
            continue
        
        # 检测标题样式
        style_name = para.style.name.lower() if para.style else ""
        
        if "heading 1" in style_name or para.style.name == "标题 1":
            markdown_lines.append(f"# {text}")
        elif "heading 2" in style_name or para.style.name == "标题 2":
            markdown_lines.append(f"## {text}")
        elif "heading 3" in style_name or para.style.name == "标题 3":
            markdown_lines.append(f"### {text}")
        elif "heading" in style_name or "标题" in para.style.name:
            level = style_name.count("heading") + style_name.count("标题") + 1
            markdown_lines.append(f"{'#' * min(level, 6)} {text}")
        else:
            # 普通段落，处理粗体和斜体
            text = _convert_formatting_to_markdown(para)
            markdown_lines.append(text)
    
    # 处理表格
    for table in doc.tables:
        markdown_lines.append("")
        markdown_lines.append(_convert_table_to_markdown(table))
        markdown_lines.append("")
    
    markdown_content = "\n".join(markdown_lines)
    
    # 保存文件
    if not output_file:
        base_name = os.path.splitext(input_file)[0]
        output_file = f"{base_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
    
    if not output_file.endswith('.md'):
        output_file += '.md'
    
    output_path = os.path.join(OUTPUT_DIR, output_file)
    output_path = _validate_path(output_path, [OUTPUT_DIR])
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown_content)
    
    return [TextContent(
        type="text",
        text=json.dumps({
            "success": True,
            "input_file": input_file,
            "output_file": output_file,
            "output_path": output_path,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "content_preview": markdown_content[:500] + "..." if len(markdown_content) > 500 else markdown_content
        }, ensure_ascii=False, indent=2)
    )]


def _convert_formatting_to_markdown(para) -> str:
    """将段落格式转换为 markdown"""
    result = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        
        # 处理格式
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        
        if run.underline:
            text = f"<u>{text}</u>"
        
        result.append(text)
    
    return "".join(result)


def _convert_table_to_markdown(table) -> str:
    """将表格转换为 markdown"""
    lines = []
    
    # 表头
    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
    lines.append("| " + " | ".join(header_cells) + " |")
    lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
    
    # 数据行
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
    
    return "\n".join(lines)


async def _handle_read_docx_content(args: Dict[str, Any]) -> List[TextContent]:
    """读取 docx 纯文本内容"""
    if not DOCX_AVAILABLE:
        return [TextContent(
            type="text",
            text=json.dumps({
                "success": False,
                "error": "python-docx 未安装"
            }, ensure_ascii=False)
        )]
    
    input_file = args["input_file"]
    input_path = os.path.join(DATA_DIR, input_file)
    input_path = _validate_path(input_path, [DATA_DIR])
    
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"文件不存在: {input_file}")
    
    doc = Document(input_path)
    
    # 提取所有文本
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # 提取表格文本
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            full_text.append(" | ".join(row_text))
    
    content = "\n\n".join(full_text)
    
    return [TextContent(
        type="text",
        text=json.dumps({
            "success": True,
            "filename": input_file,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "content": content[:2000] + "..." if len(content) > 2000 else content
        }, ensure_ascii=False, indent=2)
    )]


async def _handle_list_documents(args: Dict[str, Any]) -> List[TextContent]:
    """列出文档文件"""
    directory = args.get("directory", "data")
    extension = args.get("extension", ".docx")
    
    if directory == "data":
        target_dir = DATA_DIR
    elif directory == "plans":
        target_dir = OUTPUT_DIR
    else:
        raise ValueError(f"未知目录: {directory}")
    
    _ensure_dirs()
    
    files = []
    for filename in os.listdir(target_dir):
        if filename.endswith(extension):
            filepath = os.path.join(target_dir, filename)
            stat = os.stat(filepath)
            files.append({
                "filename": filename,
                "size": stat.st_size,
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
            })
    
    return [TextContent(
        type="text",
        text=json.dumps({
            "success": True,
            "directory": directory,
            "extension": extension,
            "count": len(files),
            "files": files
        }, ensure_ascii=False, indent=2)
    )]


async def _handle_create_sample_docx(args: Dict[str, Any]) -> List[TextContent]:
    """创建示例 docx 文件"""
    if not DOCX_AVAILABLE:
        return [TextContent(
            type="text",
            text=json.dumps({
                "success": False,
                "error": "python-docx 未安装，请运行: pip install python-docx"
            }, ensure_ascii=False)
        )]
    
    filename = args.get("filename", "sample.docx")
    title = args.get("title", "示例文档")
    
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    filepath = os.path.join(DATA_DIR, filename)
    filepath = _validate_path(filepath, [DATA_DIR])
    
    # 创建文档
    doc = Document()
    
    # 添加标题
    doc.add_heading(title, level=1)
    
    # 添加段落
    doc.add_paragraph("这是示例文档的第一段。用于测试 docx 到 markdown 的转换功能。")
    
    doc.add_heading("功能特性", level=2)
    
    # 添加带格式的段落
    para = doc.add_paragraph()
    para.add_run("支持 ")
    para.add_run("粗体").bold = True
    para.add_run(" 和 ")
    para.add_run("斜体").italic = True
    para.add_run(" 文本。")
    
    doc.add_heading("数据表格", level=2)
    
    # 添加表格
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '项目'
    hdr_cells[1].text = '数值'
    hdr_cells[2].text = '单位'
    
    # 数据行
    row1 = table.rows[1].cells
    row1[0].text = '水位'
    row1[1].text = '100.5'
    row1[2].text = 'm'
    
    row2 = table.rows[2].cells
    row2[0].text = '流量'
    row2[1].text = '500'
    row2[2].text = 'm³/s'
    
    doc.add_paragraph()
    doc.add_paragraph("文档生成时间: " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    
    # 保存
    doc.save(filepath)
    
    return [TextContent(
        type="text",
        text=json.dumps({
            "success": True,
            "filename": filename,
            "filepath": filepath,
            "title": title,
            "message": "示例文档已创建"
        }, ensure_ascii=False, indent=2)
    )]


async def main():
    """启动 MCP Server"""
    _ensure_dirs()
    
    from mcp.server.stdio import stdio_server as mcp_stdio_server
    
    async with mcp_stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream,
            write_stream,
            server.create_initialization_options()
        )


if __name__ == "__main__":
    asyncio.run(main())